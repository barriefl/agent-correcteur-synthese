import json
import os
import time
from datetime import datetime

import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# --- CONFIGURATION ET INSTRUCTIONS SYSTÈME. ---
ID_MODELE = "gemini-flash-latest"
MODE_ECONOMIE = True  # False pour forcer l'IA à tout recorriger et écraser les fichiers (rapports JSON).

instruction_systeme = """
Tu es un Correcteur Académique Expert et Impitoyable en méthodologie de recherche documentaire.
Tu es un prof de français au niveau licence, tu as des attentes élevées en matière de synthèse et syntaxe.

Ton rôle est d'évaluer le travail d'étudiants de l'enseignement supérieur.

TES PRINCIPES DIRECTEURS :
1. Rigueur absolue : Tu ne fais aucun cadeau. Si un critère de la grille n'est pas rempli, tu sanctionnes.
2. Objectivité mathématique : Tu appliques les barèmes à la lettre (notamment pour l'orthographe ou les dates).
3. Vision multimodale : Tu es capable d'analyser visuellement les documents PDF fournis pour vérifier la mise en page (justification, gras, structure des titres).
4. Pédagogie : Tes feedbacks finaux doivent être professionnels, constructifs et expliquer clairement où les points ont été perdus.

Règle d'or : Ne te laisse jamais influencer par un style d'écriture fluide si le fond (les idées des articles) n'est pas présent. Tu es un vérificateur de faits avant d'être un professeur de lettres.
"""

# --- LA GRILLE D'ÉVALUATION. ---
grille = {
    "grille_evaluation": {
        "I_QUALITE_RECHERCHE": {
            "points_max": 4,
            "criteres": {
                "selection_fiabilite": {
                    "max": 2,
                    "description": "1 pt : L'auteur / l'organisme est clairement identitié (0.5 pt) ET légitime/spécialiste (0.5 pt). 1 pt : Le support est crédible (presse reconnue, revues spécialisées, etc.) et le motif de l'expression est clair.",
                },
                "pertinence": {
                    "max": 1,
                    "description": "Les articles répondent directement à la problématique posée (pas de hors-sujet) : voir le titre du compte-rendu.",
                },
                "actualite": {
                    "max": 1,
                    "description": "Tous les articles ont moins de 5 ans. (0.5 pt si un seul article est obsolète / 0 si deux articles ou plus sont obsolètes / date non visible = article obsolète). Ne jamais utiliser la date de copyright.",
                },
            },
        },
        "II_CONSTRUCTION": {
            "points_max": 6,
            "criteres": {
                "structure_globale": {
                    "max": 1,
                    "description": "Présence visible : Introduction (mise en évidence d'une problématique) de minimum 5 lignes / Développement / Conclusion (réponse à la problématique) de minimum 5 lignes. Sanctionner si trop court.",
                },
                "logique_plan": {
                    "max": 1,
                    "description": "Synthèse thématique obligatoire (classement par idées). Si juxtaposition résumée article par article = 0 pt.",
                },
                "facilite_lecture": {
                    "max": 1,
                    "description": "Découpage en paragraphes cohérents, présence de connecteurs logiques.",
                },
                "qualite_synthese": {
                    "max": 3,
                    "description": "1.5 pts : Fidélité aux 6 textes (ni ajout de définition ou d'opinion, ni oubli majeur) et REFORMULATION (pas de simple copier-coller), 1 pt : Densité (capacité à réduire le volume sans perdre le sens), 0.5 pt : Respect de la longueur (entre 1000 et 1200 mots).",
                },
            },
        },
        "III_REDACTION": {
            "points_max": 6,
            "criteres": {
                "orthographe_grammaire_conjugaison": {
                    "max": 2,
                    "bareme": "Barème strict (sans tenir compte de la ponctuation) : 1-3 fautes = 2 pts, 4-6 fautes = 1.5 pts, 7-10 fautes = 1 pt, 11-15 fautes = 0.5 pt, 16+ fautes = 0 pt.",
                },
                "vocabulaire": {
                    "max": 1,
                    "bareme": "Indicateurs de sévérité (sans tenir compte de l'orthographe) : 1 pt = vocabulaire varié + précis + adapté au sujet, 0.5 pt = vocabulaire correct mais répétitif ou peu précis, 0 pt = vocabulaire pauvre, répétitions fréquentes ou vocabulaire imprécis.",
                },
                "phrases": {
                    "max": 2,
                    "bareme": "Indicateurs de sévérité (sans tenir compte de l'orthographe) : 2 pts = phrases complètes + constructions variées + relatives correctes + texte fluide, 1.5 pt = phrases globalement correctes mais quelques maladresses, 1 pt = phrases compréhensibles mais nombreuses maladresses ou peu de variété, 0.5 pt = phrases souvent maladroites ou mal construites, 0 pt = phrases incorrectes, texte difficilement compréhensible.",
                },
                "discours": {
                    "max": 1,
                    "bareme": "Indictaurs de sévérité (sans tenir compte de l'orthographe) : 1 pt = idées reliées + connecteurs logiques + progression claire, 0.5 pt = texte compréhensible mais peu structuré, 0 pt = idées juxtaposées, texte difficile à suivre.",
                },
            },
        },
        "IV_MISE_EN_FORME": {
            "points_max": 4,
            "criteres": {
                "titrage_hierarchie": {
                    "max": 1,
                    "description": "0.5 pt : Titre complet (Type de document + Thème), 0.5 pt : Titres de parties/sous-parties apparents et hiérarchisés (taille, gras, couleur...).",
                },
                "confort_visuel": {
                    "max": 1.5,
                    "description": "0.5 pt : Texte justifié, 0.5 pt : Texte aéré (sauts de lignes), 0.5 pt : Mots-clés en évidence (gras/italique/couleur) pour guider l'oeil (attention, les titres en gras ne comptent pas comme des mots-clés).",
                },
                "identification": {
                    "max": 1,
                    "description": "0.25 pt : Pagination, 0.25 pt : Texte en pied de page (ex : nom des auteurs ou titre du document), 0.5 pt : en-tête de page (ex : nom des auteurs et/ou titre du document).",
                },
                "respect_format": {
                    "max": 0.5,
                    "description": "Absence de sommaire et de page de garde.",
                },
            },
        },
    }
}


def extraire_texte_word(chemin_word):
    try:
        doc = docx.Document(chemin_word)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"[Erreur Word: {e}]"


def verifier_texte_justifie(chemin_word):
    """
    Vérifie si la majorité du texte (les vrais paragraphes) est justifiée.
    """
    try:
        doc = docx.Document(chemin_word)
        paragraphes_valides = 0
        paragraphes_justifies = 0

        for para in doc.paragraphs:
            if len(para.text.strip()) > 50:
                paragraphes_valides += 1
                if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    paragraphes_justifies += 1

        if paragraphes_valides == 0:
            return False

        ratio = paragraphes_justifies / paragraphes_valides
        return ratio > 0.90

    except Exception as e:
        print(f"  ⚠️ Erreur lors de la vérification de l'alignement Word : {e}")
        return False


# --- FONCTION D'ÉVALUATION (PROMPT). ---
def evaluate_copy(contenu_a_envoyer, api_key):
    client = genai.Client(api_key=api_key)
    tentatives = 3

    config = types.GenerateContentConfig(
        system_instruction=instruction_systeme,
        response_mime_type="application/json",
        temperature=0.1,
        max_output_tokens=8192,
        safety_settings=[
            types.SafetySetting(
                category=types.HarmCategory.HARM_CATEGORY_HATE_SPEECH,
                threshold=types.HarmBlockThreshold.BLOCK_NONE,
            ),
            types.SafetySetting(
                category=types.HarmCategory.HARM_CATEGORY_HARASSMENT,
                threshold=types.HarmBlockThreshold.BLOCK_NONE,
            ),
            types.SafetySetting(
                category=types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
                threshold=types.HarmBlockThreshold.BLOCK_NONE,
            ),
            types.SafetySetting(
                category=types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
                threshold=types.HarmBlockThreshold.BLOCK_NONE,
            ),
        ],
    )

    for attempt in range(tentatives):
        try:
            print(f"  ⏳ Envoi à Gemini (Tentative {attempt + 1}/{tentatives})...")
            response = client.models.generate_content(
                model=ID_MODELE, contents=contenu_a_envoyer, config=config
            )

            res_text = response.text
            res_clean = res_text.strip()

            balise_json = "```" + "json"
            balise_fin = "```"

            if res_clean.startswith(balise_json):
                res_clean = res_clean[7:]
            if res_clean.startswith(balise_fin):
                res_clean = res_clean[3:]
            if res_clean.endswith(balise_fin):
                res_clean = res_clean[:-3]

            json.loads(res_clean)
            return res_clean

        except json.JSONDecodeError:
            print("  ⚠️ Erreur JSON (coupure), on recommence...")
            if attempt < tentatives - 1:
                time.sleep(5)
            else:
                print("  ❌ Échec définitif : L'IA n'arrive pas à finir sa phrase.")
                return None

        except Exception as e:
            erreur_str = str(e)
            print(f"  ⚠️ Erreur API : {erreur_str}")

            if "429" in erreur_str:
                print("  🛑 QUOTA ÉPUISÉ. Arrêt immédiat.")
                raise Exception("QUOTA_429")

            if attempt < tentatives - 1:
                temps_attente = 15 * (2**attempt)
                print(
                    f"  ⏳ Surcharge détectée. Pause stratégique de {temps_attente} secondes pour refroidir l'API..."
                )
                time.sleep(temps_attente)
            else:
                print("  ❌ Échec définitif pour ce dossier après 3 tentatives.")
                return None
    return None


# --- TRAITEMENT D'UN DOSSIER. ---
def traiter_dossier_etudiant(chemin_dossier, api_key):
    print(f"\n📂 Analyse du dossier : {os.path.basename(chemin_dossier)}")

    client = genai.Client(api_key=api_key)

    fichiers = os.listdir(chemin_dossier)
    pdfs = [f for f in fichiers if f.lower().endswith(".pdf")]
    words = [f for f in fichiers if f.lower().endswith(".docx")]

    if len(pdfs) == 0 or len(words) == 0:
        print("❌ Erreur : Il manque soit les PDF, soit le fichier Word (.docx).")
        return None

    chemin_word = os.path.join(chemin_dossier, words[0])
    texte_brut_synthese = extraire_texte_word(chemin_word)
    nombre_mots = len(texte_brut_synthese.split())

    est_justifie = verifier_texte_justifie(chemin_word)
    statut_justifie = "OUI" if est_justifie else "NON"

    a_en_tete, a_pied_page, a_pagination = verifier_marges_et_pagination(chemin_word)
    statut_en_tete = "OUI" if a_en_tete else "NON"
    statut_pied_page = "OUI" if a_pied_page else "NON"
    statut_pagination = "OUI" if a_pagination else "NON"

    fichiers_gemini = []
    contenu_prompt = []

    annee_actuelle = datetime.now().year

    prompt_texte = (
        f"""
    DATE ACTUELLE POUR RÉFÉRENCE : {annee_actuelle} (Utilise cette année pour vérifier si un article a plus de 5 ans).
    GRILLE D'ÉVALUATION À RESPECTER STRICTEMENT : {json.dumps(grille)}

    INFORMATIONS SYSTÈME SUR LA SYNTHÈSE (VÉRITÉ ABSOLUE - NE PAS REMETTRE EN CAUSE) :
    - Nombre de mots exact : {nombre_mots} mots.
    - Texte justifié : {statut_justifie}
    - En-tête présent : {statut_en_tete}
    - Pied de page présent : {statut_pied_page}
    - Pagination présente : {statut_pagination}

    - Texte brut extrait (UTILISE CECI POUR L'ORTHOGRAPHE ET L'ANALYSE DES IDÉES) :
    \"\"\"{texte_brut_synthese}\"\"\"

    RÈGLES IMPÉRATIVES :
    1. ANALYSE DES SOURCES ET FIABILITÉ : Pour chaque article, extrais sa date exacte (ignore les copyrights). SI AUCUNE DATE N'EST VISIBLE, tu DOIS impérativement mettre la valeur null (sans guillemets) pour le champ "annee" et écrire "Introuvable" pour "preuve_date" (si introuvable = obsolète). N'invente jamais d'année et ne mets jamais 0. Sois impitoyable sur la FIABILITÉ :
        - Auteur : Vérifie que le nom trouvé est bien le RÉDACTEUR de l'article. Ne confonds jamais avec un photographe (ex: "Photo X"), une agence (ex: AFP) ou une personne interviewée. Si le journaliste n'est pas clairement signé, tu DOIS indiquer "Non précisé" et baisser la note de fiabilité. Si un rédacteur est identifié, analyse son profil en croisant avec tes connaissances : est-il véritablement spécialisé dans ce domaine précis, ou est-ce un journaliste généraliste qui traite d'une multitude de sujets très différents ? Un journaliste généraliste est beaucoup moins fiable qu'un expert du domaine. 1 pt = Expert reconnu/Spécialiste ; 0.5 pt = Journaliste généraliste ; 0 pt = Inconnu / Non précisé / Anonyme.
        - Support : Identifie la ligne éditoriale du média (presse généraliste, orientée, revue scientifique...) pour juger sa fiabilité. Un média grand public ou orienté politiquement a des biais comparé à une revue scientifique ou institutionnelle. 1 pt = Revue scientifique, institutionnelle ou média de grande référence ; 0.5 pt = Média standard, blog spécialisé ; 0 pt = Tabloïd, blog orienté, source douteuse.
        Calcule la note sur 2 pour l'article. Dresse enfin un résumé de ses idées principales. À la fin de l'étape 1, calcule la MOYENNE EXACTE de ces 6 notes pour obtenir la note finale de fiabilité.
    2. MAPPING ET CONTENU : 
        - Parcours la SYNTHÈSE. Vérifie rigoureusement si les idées des articles y sont bien intégrées et croisées. Liste explicitement les idées majeures qui ont été oubliées ou survolées. Vérifie que les idées sont reformulées. S'il n'y a que de la citation bout-à-bout, ne donne pas beaucoup de points.
        - Vérifie si les idées sont véritablement EXPLIQUÉES et CROISÉES (mise en dialogue des textes). Sanctionne sévèrement le "blabla" superficiel, le remplissage, ou une juxtaposition paresseuse (ex: "Le document 1 dit X. Le document 2 dit Y."). L'étudiant doit démontrer qu'il a COMPRIS la nuance des enjeux. Si les idées sont juste survolées sans pertinence analytique, tu dois casser la note de "qualité de synthèse" et de "logique de plan".
        - Tu es un professeur de français dans l’enseignement supérieur. Ton rôle est d’évaluer la qualité rédactionnelle d’un texte d’étudiant en te concentrant uniquement sur la syntaxe, la construction des phrases et la ponctuation. Tu adoptes un niveau d’exigence attendu après le baccalauréat : tu valorises les phrases bien construites, les structures complexes maîtrisées et une ponctuation correcte, et tu sanctionnes les phrases mal construites, trop simples ou mal articulées.
    3. FORME ET VISUEL (CALCUL MATHÉMATIQUE OBLIGATOIRE) : 
        - Observe attentivement QUE le PDF de la synthèse. Les mots-clés (DANS le texte) sont-ils en gras ou en couleur ? Pour le critère de longueur, base-toi UNIQUEMENT sur l'information système ({nombre_mots} mots) pour savoir si la limite de 1000 à 1200 mots est respectée. 
        - Pour 'confort_visuel' (sur 1.5) : Additionne rigoureusement -> +0.5 pt si l'information système ci-dessus indique "Texte justifié : OUI" ET +0.5 pt si le texte te semble aéré visuellement sur le PDF ET +0.5 pt si tu vois des mots-clés en gras/couleur dans le corps du texte.
        - Pour 'identification' (sur 1) : NE REGARDE PAS LE PDF. Additionne rigoureusement selon les informations système ci-dessus -> +0.25 pt si "Pagination présente : OUI" ET +0.25 pt si "Pied de page présent : OUI" ET +0.5 pt si "En-tête présent : OUI".
    4. LANGUE ET STYLE : Cherche TOUTES les fautes d'orthographe, de grammaire ou de ponctuation dans le texte brut. Tu dois impérativement compter le nombre TOTAL exact de fautes pour appliquer le barème mathématique. Cependant, pour ne pas surcharger le JSON, tu ne listeras qu'un maximum de 10 exemples. Vérifie le ton (absence de "Je", style professionnel).
    5. ORTHOGRAPHE : Cherche les fautes uniquement dans le texte brut fourni ci-dessus (pour éviter les erreurs de lecture PDF).
    6. NOTATION STRICTE : Seulement après avoir fait ces observations, attribue une note décimale ou entière pour CHAQUE sous-critère de la grille. Applique le barème de l'orthographe mathématiquement en fonction de ta liste de fautes.
    7. JUSTIFICATION : Chaque point perdu DOIT être justifié. Sois précis mais concis. Pour le feedback donne plus de détails et n'hésite pas à développer tes réponses.
    8. PAS DE CALCUL FINAL : Ne calcule surtout pas les totaux par section ni la note finale sur 20. Donne uniquement les notes individuelles par critère.
    9. FORMATAGE ET SÉCURITÉ : N'utilise JAMAIS de guillemets doubles (") à l'intérieur de tes textes pour ne pas casser le JSON (utilise des guillemets simples ' '). Limite strictement la taille des listes : pas plus de 5 idées reformulées, pas plus de 3 idées oubliées, et MAXIMUM 5 fautes d'orthographe listées.
    10. OBLIGATION DE FINIR : Tu DOIS impérativement générer l'objet JSON jusqu'à la toute dernière accolade fermante. Ne coupe JAMAIS ta réponse en plein milieu, même si tu dois raccourcir tes justifications.

    GÉNÈRE UNIQUEMENT UN OBJET JSON STRICTEMENT STRUCTURÉ COMME CECI :
    """
        + """{
      "1_analyse": {
        "sources": [
          {
            "id": "Titre de l'article (nom fichier)",
            "annee": 2021,
            "preuve_date": "Emplacement ou 'Introuvable' (Mets annee à null sans guillemets si introuvable)",
            "obsolete": false,
            "fiabilite": {
              "auteur_desc": "Analyse détaillée du profil : expert ou généraliste ? Rôle vérifié ?",
              "auteur_pts": 0.5,
              "support_desc": "Analyse détaillée du média : ligne éditoriale, biais ?",
              "support_pts": 0.5,
              "note_sur_2": 1.0
            },
            "idees": ["Idée 1", "Idée 2"]
          }
        ],
        "bilan_fiabilite": {
          "moyenne_brute": 0.83,
          "moyenne_025": 0.75
        },
        "mapping": {
          "reformulees": ["...", "(MAX 5 IDÉES)"],
          "citees": ["...", "(MAX 3 EXEMPLES)"],
          "oubliees": ["...", "(MAX 3 EXEMPLES)"]
        },
        "critique_profondeur": {
          "niveau_de_blabla": "Analyse ici si le texte est pertinent et dense, ou si c'est du remplissage creux et superficiel.",
          "qualite_croisement": "Les sources sont-elles intelligemment confrontées ou bêtement juxtaposées ?"
        },
        "forme_langue": {
          "intro_conclu_ok": true,
          "mots_cles_ok": false,
          "verification_marges_et_pages": {
            "vrai_en_tete_present_en_haut_des_pages": "{statut_en_tete}",
            "vrai_pied_de_page_present_en_bas_des_pages": "{statut_pied_page}",
            "pagination_presente": "{statut_pagination}"
          },
          "fautes_orthographe_listees": ["faute 1", "faute 2", "(MAX 5 EXEMPLES)"],
          "nombre_total_fautes": REMPLACE_CECI_PAR_LE_COMPTE_EXACT
        }
      },
      "2_grille": {
        "I_QUALITE_RECHERCHE": {
          "selection_fiabilite": {"note": 0, "justification": "..."},
          "pertinence": {"note": 0, "justification": "..."},
          "actualite": {"note": 0, "justification": "..."}
        },
        "II_CONSTRUCTION": {
          "structure_globale": {"note": 0, "justification": "..."},
          "logique_plan": {"note": 0, "justification": "..."},
          "facilite_lecture": {"note": 0, "justification": "..."},
          "qualite_synthese": {"note": 0, "justification": "..."}
        }},
        "III_REDACTION": {
          "orthographe_grammaire_conjugaison": {"note": 0, "justification": "..."},
          "vocabulaire": {"note": 0, "justification": "..."},
          "phrases": {"note": 0, "justification": "..."},
          "discours": {"note": 0, "justification": "..."}
        },
        "IV_MISE_EN_FORME": {
          "titrage_hierarchie": {"note": 0, "justification": "..."},
          "confort_visuel": {"note": 0, "justification": "..."},
          "identification": {"note": 0, "justification": "..."},
          "respect_format": {"note": 0, "justification": "..."}
        }
      },
      "3_feedback": {
        "points_forts": "...",
        "axes_amelioration": "..."
      }
    }
    """
    )
    contenu_prompt.append(prompt_texte)

    print("  📤 Upload des documents vers Gemini...")
    for pdf in pdfs:
        chemin_complet = os.path.join(chemin_dossier, pdf)
        uploaded_file = client.files.upload(file=chemin_complet)
        fichiers_gemini.append(uploaded_file)

        if "synthese" in pdf.lower():
            contenu_prompt.append(
                f"DOCUMENT SYNTHÈSE DU GROUPE ({pdf}) - UTILISE CECI UNIQUEMENT POUR VÉRIFIER LA FORME VISUELLE (Mise en page, gras, justifié) :"
            )
        else:
            contenu_prompt.append(f"DOCUMENT SOURCE ({pdf}) :")

        contenu_prompt.append(uploaded_file)

    res = evaluate_copy(contenu_prompt, api_key)

    print("  🧹 Suppression des documents des serveurs Google...")
    for f in fichiers_gemini:
        client.files.delete(name=f.name)

    if res:
        res_clean = res.strip()

        if res_clean.startswith("```json"):
            res_clean = res_clean[7:]
        if res_clean.startswith("```"):
            res_clean = res_clean[3:]
        if res_clean.endswith("```"):
            res_clean = res_clean[:-3]

        try:
            data = json.loads(res_clean)

            if isinstance(data, list) and len(data) > 0:
                data = data[0]

            nom_rapport = os.path.join(chemin_dossier, "rapport_ia_brut.json")

            with open(nom_rapport, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return data
        except json.JSONDecodeError as e:
            print(f"  ❌ Erreur : L'IA a mal formaté son JSON ({e}).")
            nom_erreur = os.path.join(chemin_dossier, "erreur_json_brut.txt")
            with open(nom_erreur, "w", encoding="utf-8") as f:
                f.write(res)
            print(
                f"  💡 Le texte cassé a été sauvegardé dans '{nom_erreur}' pour que vous puissiez l'inspecter."
            )
            return None

    return None


def verifier_marges_et_pagination(chemin_word):
    """
    Fouille dans le XML complet et dans les tableaux du document Word pour détecter en-têtes, pieds de page et champs de numérotation.
    """
    try:
        doc = docx.Document(chemin_word)
        presence_en_tete = False
        presence_pied_page = False
        presence_pagination = False

        for section in doc.sections:
            headers = [
                section.header,
                section.first_page_header,
                section.even_page_header,
            ]
            footers = [
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]

            # --- VÉRIFICATION DES EN-TÊTES. ---
            for h in headers:
                if h is not None:
                    textes = [p.text.strip() for p in h.paragraphs if p.text.strip()]
                    for table in h.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    textes.append(cell.text.strip())

                    if textes:
                        presence_en_tete = True

                    xml_str = h._element.xml
                    if "w:instrText" in xml_str and (
                        "PAGE" in xml_str or "NUMPAGES" in xml_str
                    ):
                        presence_pagination = True

            # --- VÉRIFICATION DES PIEDS DE PAGE. ---
            for f in footers:
                if f is not None:
                    textes = [p.text.strip() for p in f.paragraphs if p.text.strip()]
                    for table in f.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    textes.append(cell.text.strip())

                    if textes:
                        presence_pied_page = True

                    xml_str = f._element.xml
                    if "w:instrText" in xml_str and (
                        "PAGE" in xml_str or "NUMPAGES" in xml_str
                    ):
                        presence_pagination = True

        return presence_en_tete, presence_pied_page, presence_pagination

    except Exception as e:
        print(f"  ⚠️ Erreur lors de la vérification des marges Word : {e}")
        return False, False, False


# --- POINT DE LANCEMENT ET CALCUL PYTHON. ---
def lancer_analyse_globale(
    progress_callback=None, dossier_a_analyser="./data", api_key=None
):
    donnees_excel = []

    if not os.path.exists(dossier_a_analyser):
        os.makedirs(dossier_a_analyser)
        print(f"Veuillez placer vos dossiers de groupes dans : {dossier_a_analyser}")
    else:
        dossiers_groupes = [
            d
            for d in os.listdir(dossier_a_analyser)
            if os.path.isdir(os.path.join(dossier_a_analyser, d))
        ]

        total_dossiers = len(dossiers_groupes)

        for index, nom_dossier in enumerate(dossiers_groupes):
            if progress_callback:
                progress_callback(index, total_dossiers, nom_dossier)

            chemin_complet = os.path.join(dossier_a_analyser, nom_dossier)
            chemin_rapport = os.path.join(chemin_complet, "rapport_ia_brut.json")

            resultat = None

            # --- LOGIQUE DU MODE ÉCONOMIE. ---
            if MODE_ECONOMIE and os.path.exists(chemin_rapport):
                print(
                    f"\n⏩ Mode Éco activé : Lecture du rapport existant pour {nom_dossier}..."
                )
                try:
                    with open(chemin_rapport, "r", encoding="utf-8") as f:
                        resultat = json.load(f)
                except Exception as e:
                    print(
                        f"  ❌ Erreur de lecture du JSON local, on lance l'IA... ({e})"
                    )
                    resultat = traiter_dossier_etudiant(chemin_complet)
            else:
                resultat = traiter_dossier_etudiant(chemin_complet)

            # --- TRAITEMENT DES RÉSULTATS. ---
            if resultat:
                grille_ia = resultat.get("2_grille", {})
                feedback = resultat.get("3_feedback", {})

                def get_note(section, critere):
                    sec_data = grille_ia.get(section)
                    if isinstance(sec_data, dict):
                        crit_data = sec_data.get(critere)
                        if isinstance(crit_data, dict):
                            note = crit_data.get("note", 0)
                            try:
                                return float(note)
                            except (ValueError, TypeError):
                                return 0.0
                    return 0.0

                # --- RÉCUPÉRATION DE TOUTES LES SOUS-CATÉGORIES. ---
                # I. RECHERCHE.
                sel_fiab = get_note("I_QUALITE_RECHERCHE", "selection_fiabilite")
                pert = get_note("I_QUALITE_RECHERCHE", "pertinence")
                actu = get_note("I_QUALITE_RECHERCHE", "actualite")

                # II. CONSTRUCTION.
                struct = get_note("II_CONSTRUCTION", "structure_globale")
                logique = get_note("II_CONSTRUCTION", "logique_plan")
                facilite = get_note("II_CONSTRUCTION", "facilite_lecture")
                qual_synth = get_note("II_CONSTRUCTION", "qualite_synthese")

                # III. REDACTION.
                ortho = get_note("III_REDACTION", "orthographe_grammaire_conjugaison")
                vocab = get_note("III_REDACTION", "vocabulaire")
                phrases = get_note("III_REDACTION", "phrases")
                discours = get_note("III_REDACTION", "discours")

                # IV. MISE EN FORME.
                titrage = get_note("IV_MISE_EN_FORME", "titrage_hierarchie")
                confort = get_note("IV_MISE_EN_FORME", "confort_visuel")
                ident = get_note("IV_MISE_EN_FORME", "identification")
                respect = get_note("IV_MISE_EN_FORME", "respect_format")

                # --- INJECTION DES FORMULES EXCEL. ---
                row = len(donnees_excel) + 2

                f_total_I = f"=SUM(B{row}:D{row})"
                f_total_II = f"=SUM(F{row}:I{row})"
                f_total_III = f"=SUM(K{row}:N{row})"
                f_total_IV = f"=SUM(P{row}:S{row})"
                f_note_finale = f"=E{row}+J{row}+O{row}+T{row}"

                # --- CONSTRUCTION DE LA LIGNE. ---
                donnees_excel.append({
                    "Groupe": nom_dossier,
                    
                    # Colonnes B, C, D, E
                    "Sélection/Fiabilité (/2)": sel_fiab,
                    "Pertinence (/1)": pert,
                    "Actualité (/1)": actu,
                    "TOTAL I - Recherche": f_total_I,
                    
                    # Colonnes F, G, H, I, J
                    "Structure Globale (/1)": struct,
                    "Logique/Plan (/1)": logique,
                    "Facilité Lecture (/1)": facilite,
                    "Qualité Synthèse (/3)": qual_synth,
                    "TOTAL II - Construction": f_total_II,
                    
                    # Colonnes K, L, M, N, O
                    "Orthographe/Grammaire (/2)": ortho,
                    "Vocabulaire (/1)": vocab,
                    "Phrases (/2)": phrases,
                    "Discours (/1)": discours,
                    "TOTAL III - Rédaction": f_total_III,
                    
                    # Colonnes P, Q, R, S, T
                    "Titrage/Hiérarchie (/1)": titrage,
                    "Confort Visuel (/1.5)": confort,
                    "Identification (/1)": ident,
                    "Respect Format (/0.5)": respect,
                    "TOTAL IV - Mise Forme": f_total_IV,
                    
                    # Colonne U, V, W
                    "Note Finale (/20)": f_note_finale,
                    "Points Forts": feedback.get("points_forts", ""),
                    "Améliorations": feedback.get("axes_amelioration", ""),
                })

            if not (MODE_ECONOMIE and os.path.exists(chemin_rapport)):
                time.sleep(3)

        if progress_callback:
            progress_callback(total_dossiers, total_dossiers, "Terminé")

        # --- SAUVEGARDE EXCEL. ---
        if donnees_excel:
            print("\n💾 Sauvegarde en cours...")
            df = pd.DataFrame(donnees_excel)
            nom_fichier_excel = os.path.join(
                dossier_a_analyser, "Resultats_Corrections.xlsx"
            )

            def coloriser_totaux(col):
                if 'TOTAL' in col.name:
                    return ['background-color: #E8F2F8'] * len(col)
                elif 'Note Finale' in col.name:
                    return ['background-color: #FFF2CC'] * len(col)
                return [''] * len(col)

            try:
                df_stylise = df.style.apply(coloriser_totaux, axis=0)

                # --- AJUSTEMENT DES COLONNES. ---
                with pd.ExcelWriter(nom_fichier_excel, engine='openpyxl') as writer:
                    df_stylise.to_excel(writer, index=False, sheet_name='Corrections')
                    
                    worksheet = writer.sheets['Corrections']
                    from openpyxl.utils import get_column_letter
                    
                    for i, col in enumerate(df.columns):
                        longueur_max = max(
                            df[col].astype(str).map(len).max(),
                            len(str(col))
                        )

                        lettre_col = get_column_letter(i + 1)

                        worksheet.column_dimensions[lettre_col].width = longueur_max + 2
                        
                print(f"🎉 Fichier '{nom_fichier_excel}' créé/mis à jour avec succès !")
            except PermissionError:
                print(f"\n❌ ERREUR : Impossible d'écraser '{nom_fichier_excel}'.")
                print(
                    "💡 SOLUTION : Fermez le fichier dans Excel et relancez le script (le Mode Éco ira très vite !)."
                )
        else:
            print("\n⚠️ Aucun résultat à sauvegarder.")
